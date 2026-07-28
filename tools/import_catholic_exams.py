#!/usr/bin/env python3

from __future__ import annotations

import argparse
import json
import re
import shutil
import sys
from collections import Counter
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

try:
    from docx import Document
except ImportError:
    Document = None


ROOT = Path(__file__).resolve().parents[1]
PACK = ROOT / "packs" / "catholic-high-school-entrance-exams"
DATA_DIR = PACK / "data"
IMPORTS_DIR = ROOT / "imports"
REPORTS_DIR = IMPORTS_DIR / "reports"
OVERRIDES_PATH = IMPORTS_DIR / "item_overrides.json"


FAMILY_CONFIG: dict[str, dict[str, Any]] = {
    "hspt": {
        "display": "HSPT",
        "prefix": "HSPT",
        "exam_total": 298,
        "part_totals": {"a": 75, "b": 75, "c": 74, "d": 74},
        "sections": [
            {
                "source": "Verbal Skills",
                "slug": "verbal_skills",
                "prefix": "V",
                "count": 60,
            },
            {
                "source": "Quantitative Skills",
                "slug": "quantitative_skills",
                "prefix": "Q",
                "count": 52,
            },
            {
                "source": "Reading Comprehension",
                "slug": "reading_comprehension",
                "prefix": "R",
                "count": 62,
            },
            {
                "source": "Mathematics",
                "slug": "mathematics",
                "prefix": "M",
                "count": 64,
            },
            {
                "source": "Language",
                "slug": "language",
                "prefix": "L",
                "count": 60,
            },
        ],
    },
    "tachs": {
        "display": "TACHS",
        "prefix": "TACHS",
        "exam_total": 200,
        "part_totals": {"a": 50, "b": 50, "c": 50, "d": 50},
        "sections": [
            {
                "source": "Reading",
                "slug": "reading",
                "prefix": "R",
                "count": 50,
            },
            {
                "source": "Written Expression",
                "slug": "written_expression",
                "prefix": "WE",
                "count": 50,
            },
            {
                "source": "Mathematics",
                "slug": "mathematics",
                "prefix": "M",
                "count": 50,
            },
            {
                "source": "Ability",
                "slug": "ability",
                "prefix": "A",
                "count": 50,
            },
        ],
    },
    "coop": {
        "display": "COOP",
        "prefix": "COOP",
        "exam_total": 200,
        "part_totals": {"a": 50, "b": 50, "c": 50, "d": 50},
        "sections": [
            {
                "source": "Mathematics",
                "slug": "mathematics",
                "prefix": "M",
                "count": 40,
            },
            {
                "source": "Reading Comprehension",
                "slug": "reading_comprehension",
                "prefix": "R",
                "count": 40,
            },
            {
                "source": "Vocabulary",
                "slug": "vocabulary",
                "prefix": "V",
                "count": 40,
            },
            {
                "source": "Grammar",
                "slug": "grammar",
                "prefix": "G",
                "count": 40,
            },
            {
                "source": "Spelling",
                "slug": "spelling",
                "prefix": "S",
                "count": 40,
            },
        ],
    },
}


QUESTION_ID_RE = re.compile(
    r"^(HSPT-(?:V|Q|R|M|L)\d+-\d{3}|"
    r"TACHS-(?:R|WE|M|A)\d+-\d{3}|"
    r"COOP-(?:M|R|V|G|S)\d+-\d{3})$"
)

STIMULUS_ID_RE = re.compile(
    r"^(HSPT-(?:R|L)\d+-S\d{2}|"
    r"TACHS-(?:R|WE)\d+-S\d{2}|"
    r"COOP-(?:R|G)\d+-S\d{2})$"
)

ANSWER_RE = re.compile(
    r"^(?P<id>[A-Z]+-(?:[A-Z]+)?\d+-\d{3})"
    r"\s*[—–-]\s*Correct:\s*(?P<letter>[A-D])"
    r"\s*[—–-]\s*Correct Answer:\s*(?P<answer>.*?)"
    r"\s*[—–-]\s*Explanation:\s*(?P<explanation>.*)$",
    re.IGNORECASE,
)

FILE_RE = re.compile(
    r"(?P<family>hspt|tachs|coop)_full_exam_"
    r"(?P<exam>\d{1,2})_?part_?(?P<part>[abcd])",
    re.IGNORECASE,
)


@dataclass
class Stimulus:
    stimulus_id: str
    section: str
    stimulus_type: str
    title: str
    text: str


@dataclass
class Question:
    qid: str
    section: str
    category: str
    skill: str
    stimulus_id: str
    prompt: str
    choices: dict[str, str]
    correct: str = ""
    correct_answer: str = ""
    explanation: str = ""


def clean_text(value: str) -> str:
    value = str(value or "")
    value = value.replace("\u00a0", " ")
    value = value.replace("\ufeff", "")
    value = value.replace("−", "−")
    value = value.replace("\r\n", "\n").replace("\r", "\n")
    value = re.sub(r"[ \t]+\n", "\n", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()


def read_docx(path: Path) -> str:
    if Document is None:
        raise RuntimeError(
            "python-docx is required for DOCX files. "
            "Install it with: python3 -m pip install python-docx"
        )

    document = Document(path)
    lines: list[str] = []

    for paragraph in document.paragraphs:
        text = clean_text(paragraph.text)
        lines.append(text)

    for table in document.tables:
        for row in table.rows:
            values = [clean_text(cell.text) for cell in row.cells]
            lines.append(" | ".join(values))

    return "\n".join(lines)


def read_source(path: Path) -> str:
    if path.suffix.lower() == ".docx":
        return read_docx(path)

    return path.read_text(encoding="utf-8-sig")


def normalize_lines(text: str) -> list[str]:
    raw_lines = text.splitlines()
    lines: list[str] = []

    for raw in raw_lines:
        line = clean_text(raw)

        # Accept occasional "ID: " prefixes generated before item IDs.
        line = re.sub(
            r"^ID:\s*(?=(?:HSPT|TACHS|COOP)-)",
            "",
            line,
            flags=re.IGNORECASE,
        )

        if line:
            lines.append(line)

    return lines


def discover_sources(family: str) -> dict[int, dict[str, Path]]:
    folder = IMPORTS_DIR / family
    result: dict[int, dict[str, Path]] = {}

    if not folder.exists():
        return result

    candidates = sorted(
        list(folder.rglob("*.docx")) +
        list(folder.rglob("*.txt"))
    )

    for path in candidates:
        match = FILE_RE.search(path.stem)

        if not match:
            continue

        if match.group("family").lower() != family:
            continue

        exam_no = int(match.group("exam"))
        part = match.group("part").lower()

        result.setdefault(exam_no, {})

        current = result[exam_no].get(part)

        if current is None:
            result[exam_no][part] = path
        elif current.suffix.lower() == ".txt" and path.suffix.lower() == ".docx":
            result[exam_no][part] = path

    return result


def parse_stimulus(lines: list[str], start: int) -> tuple[Stimulus, int]:
    data: dict[str, str] = {
        "Stimulus ID": "",
        "Section": "",
        "Stimulus Type": "",
        "Title": "",
        "Text": "",
    }

    index = start + 1
    current_field: str | None = None
    text_lines: list[str] = []

    while index < len(lines):
        line = lines[index]

        if line == "STIMULUS END":
            break

        matched_field = False

        for field in data:
            prefix = f"{field}:"

            if line.startswith(prefix):
                if current_field == "Text":
                    data["Text"] = "\n".join(text_lines).strip()
                    text_lines = []

                current_field = field
                value = line[len(prefix):].strip()

                if field == "Text":
                    if value:
                        text_lines.append(value)
                else:
                    data[field] = value

                matched_field = True
                break

        if not matched_field and current_field == "Text":
            text_lines.append(line)
        elif not matched_field and current_field in {"Title", "Stimulus Type"}:
            data[current_field] = (
                f'{data[current_field]} {line}'
            ).strip()

        index += 1

    if current_field == "Text":
        data["Text"] = "\n".join(text_lines).strip()

    stimulus_id = data["Stimulus ID"]

    if not STIMULUS_ID_RE.match(stimulus_id):
        raise ValueError(f"Invalid Stimulus ID: {stimulus_id!r}")

    return (
        Stimulus(
            stimulus_id=stimulus_id,
            section=data["Section"],
            stimulus_type=data["Stimulus Type"],
            title=data["Title"],
            text=data["Text"],
        ),
        index + 1,
    )


def collect_field(
    lines: list[str],
    index: int,
    field_name: str,
    stop_prefixes: tuple[str, ...],
) -> tuple[str, int]:
    prefix = f"{field_name}:"
    first = lines[index][len(prefix):].strip()
    values = [first] if first else []
    index += 1

    while index < len(lines):
        line = lines[index]

        if any(line.startswith(stop) for stop in stop_prefixes):
            break

        if QUESTION_ID_RE.match(line):
            break

        if line in {
            "ANSWER KEY + EXPLANATIONS",
            "STIMULUS START",
            "STIMULUS END",
        }:
            break

        values.append(line)
        index += 1

    return clean_text("\n".join(values)), index



def is_structural_heading(value: str) -> bool:
    value = clean_text(value)

    if re.match(
        r"^SECTION\s+\d+\s*[-–—:]\s*.+$",
        value,
        flags=re.IGNORECASE,
    ):
        return True

    if (
        value == value.upper()
        and re.match(
            r"^[A-Z0-9][A-Z0-9 &/()'’.,\-]*\s+ANSWERS$",
            value,
        )
    ):
        return True

    return False


def parse_question(lines: list[str], start: int) -> tuple[Question, int]:
    qid = lines[start]
    index = start + 1

    fields: dict[str, str] = {
        "Section": "",
        "Category": "",
        "Skill": "",
        "Stimulus ID": "None",
        "Prompt": "",
    }
    choices: dict[str, str] = {}

    while index < len(lines):
        line = lines[index]

        if QUESTION_ID_RE.match(line):
            break

        if line in {
            "STIMULUS START",
            "ANSWER KEY + EXPLANATIONS",
        }:
            break

        if line.startswith("Section:"):
            fields["Section"] = line.split(":", 1)[1].strip()

        elif line.startswith("Type:"):
            pass

        elif line.startswith("Category:"):
            fields["Category"] = line.split(":", 1)[1].strip()

        elif line.startswith("Skill:"):
            fields["Skill"] = line.split(":", 1)[1].strip()

        elif line.startswith("Stimulus ID:"):
            fields["Stimulus ID"] = line.split(":", 1)[1].strip()

        elif line.startswith("Prompt:"):
            fields["Prompt"], index = collect_field(
                lines,
                index,
                "Prompt",
                ("A)", "B)", "C)", "D)"),
            )
            continue

        elif re.match(r"^[A-D]\)", line):
            letter = line[0]
            value = line[2:].strip()
            index += 1
            continuation: list[str] = [value] if value else []

            while index < len(lines):
                next_line = lines[index]

                if is_structural_heading(next_line):
                    break

                if re.match(r"^[A-D]\)", next_line):
                    break

                if QUESTION_ID_RE.match(next_line):
                    break

                if next_line in {
                    "STIMULUS START",
                    "ANSWER KEY + EXPLANATIONS",
                }:
                    break

                if any(
                    next_line.startswith(prefix)
                    for prefix in (
                        "Section:",
                        "Type:",
                        "Category:",
                        "Skill:",
                        "Stimulus ID:",
                        "Prompt:",
                    )
                ):
                    break

                continuation.append(next_line)
                index += 1

            choices[letter] = clean_text("\n".join(continuation))
            continue

        index += 1

    return (
        Question(
            qid=qid,
            section=fields["Section"],
            category=fields["Category"],
            skill=fields["Skill"],
            stimulus_id=fields["Stimulus ID"] or "None",
            prompt=fields["Prompt"],
            choices=choices,
        ),
        index,
    )


def merge_wrapped_answer_lines(lines: list[str]) -> list[str]:
    merged: list[str] = []
    current = ""

    for line in lines:
        if is_structural_heading(line):
            continue

        if re.match(
            r"^[A-Z]+-(?:[A-Z]+)?\d+-\d{3}\s*[—–-]\s*Correct:",
            line,
        ):
            if current:
                merged.append(current.strip())
            current = line
        elif current:
            current += " " + line

    if current:
        merged.append(current.strip())

    return merged


def parse_part(path: Path) -> tuple[list[Question], dict[str, Stimulus], list[str]]:
    lines = normalize_lines(read_source(path))
    questions: list[Question] = []
    stimuli: dict[str, Stimulus] = {}
    warnings: list[str] = []

    try:
        answer_index = lines.index("ANSWER KEY + EXPLANATIONS")
    except ValueError:
        answer_index = len(lines)
        warnings.append("Missing ANSWER KEY + EXPLANATIONS heading.")

    question_lines = lines[:answer_index]
    answer_lines = lines[answer_index + 1:]

    index = 0

    while index < len(question_lines):
        line = question_lines[index]

        if line == "STIMULUS START":
            try:
                stimulus, index = parse_stimulus(question_lines, index)
                stimuli[stimulus.stimulus_id] = stimulus
            except Exception as exc:
                warnings.append(f"Stimulus parse error near line {index + 1}: {exc}")
                index += 1
            continue

        if QUESTION_ID_RE.match(line):
            try:
                question, index = parse_question(question_lines, index)
                questions.append(question)
            except Exception as exc:
                warnings.append(f"Question parse error for {line}: {exc}")
                index += 1
            continue

        index += 1

    answers: dict[str, tuple[str, str, str]] = {}

    for line in merge_wrapped_answer_lines(answer_lines):
        match = ANSWER_RE.match(line)

        if not match:
            warnings.append(f"Unparsed answer-key line: {line[:160]}")
            continue

        answers[match.group("id").upper()] = (
            match.group("letter").upper(),
            clean_text(match.group("answer")),
            clean_text(match.group("explanation")),
        )

    for question in questions:
        answer = answers.get(question.qid.upper())

        if answer:
            question.correct = answer[0]
            question.correct_answer = answer[1]
            question.explanation = answer[2]

    return questions, stimuli, warnings


def load_overrides() -> dict[str, Any]:
    if not OVERRIDES_PATH.exists():
        return {
            "replacements": {},
            "answerCorrections": {},
            "explanationCorrections": {},
        }

    return json.loads(OVERRIDES_PATH.read_text(encoding="utf-8"))


def question_from_override(qid: str, value: dict[str, Any]) -> Question:
    return Question(
        qid=qid,
        section=str(value.get("section", "")),
        category=str(value.get("category", value.get("section", ""))),
        skill=str(value.get("skill", "")),
        stimulus_id=str(value.get("stimulusId", "None")),
        prompt=str(value.get("prompt", "")),
        choices={
            letter: str(value.get("choices", {}).get(letter, ""))
            for letter in "ABCD"
        },
        correct=str(value.get("correct", "")).upper(),
        correct_answer=str(value.get("correctAnswer", "")),
        explanation=str(value.get("explanation", "")),
    )


def apply_overrides(questions: list[Question]) -> list[str]:
    overrides = load_overrides()
    notes: list[str] = []

    replacements = overrides.get("replacements", {})
    answer_corrections = overrides.get("answerCorrections", {})
    explanation_corrections = overrides.get(
        "explanationCorrections",
        {},
    )

    by_id = {question.qid: question for question in questions}

    for qid, value in replacements.items():
        if qid in by_id:
            replacement = question_from_override(qid, value)
            index = questions.index(by_id[qid])
            questions[index] = replacement
            by_id[qid] = replacement
            notes.append(f"Applied full replacement: {qid}")

    for qid, value in answer_corrections.items():
        question = by_id.get(qid)

        if not question:
            continue

        if isinstance(value, str):
            question.correct = value.upper()
        elif isinstance(value, dict):
            if value.get("correct"):
                question.correct = str(value["correct"]).upper()
            if value.get("correctAnswer"):
                question.correct_answer = str(value["correctAnswer"])

        notes.append(f"Applied answer correction: {qid}")

    for qid, explanation in explanation_corrections.items():
        question = by_id.get(qid)

        if question:
            question.explanation = str(explanation)
            notes.append(f"Applied explanation correction: {qid}")

    return notes


def expected_ids(family: str, exam_no: int, section: dict[str, Any]) -> list[str]:
    prefix = FAMILY_CONFIG[family]["prefix"]
    item_prefix = section["prefix"]

    return [
        f"{prefix}-{item_prefix}{exam_no}-{number:03d}"
        for number in range(1, section["count"] + 1)
    ]


def normalize_option_text(value: str) -> str:
    return re.sub(
        r"\s+",
        " ",
        clean_text(value),
    ).strip()


def validate_exam(
    family: str,
    exam_no: int,
    questions: list[Question],
    stimuli: dict[str, Stimulus],
    parts: dict[str, Path],
) -> tuple[list[str], list[str]]:
    errors: list[str] = []
    warnings: list[str] = []
    config = FAMILY_CONFIG[family]

    ids = [question.qid for question in questions]
    id_counts = Counter(ids)

    for qid, count in id_counts.items():
        if count > 1:
            errors.append(f"Duplicate question ID: {qid} ({count} occurrences)")

    if len(questions) != config["exam_total"]:
        errors.append(
            f'Expected {config["exam_total"]} questions, found {len(questions)}.'
        )

    for part, path in parts.items():
        part_questions, _, _ = parse_part(path)
        expected_part_total = config["part_totals"][part]

        if len(part_questions) != expected_part_total:
            errors.append(
                f"{path.name}: expected {expected_part_total} questions, "
                f"found {len(part_questions)}."
            )

    section_names = {
        section["source"]: section
        for section in config["sections"]
    }

    for section in config["sections"]:
        section_questions = [
            question
            for question in questions
            if question.section == section["source"]
        ]

        expected = expected_ids(family, exam_no, section)
        actual = [question.qid for question in section_questions]

        if len(section_questions) != section["count"]:
            errors.append(
                f'{section["source"]}: expected {section["count"]} questions, '
                f"found {len(section_questions)}."
            )

        missing = sorted(set(expected) - set(actual))
        extra = sorted(set(actual) - set(expected))

        if missing:
            errors.append(
                f'{section["source"]}: missing IDs: {", ".join(missing)}'
            )

        if extra:
            errors.append(
                f'{section["source"]}: unexpected IDs: {", ".join(extra)}'
            )

    for question in questions:
        if question.section not in section_names:
            errors.append(
                f"{question.qid}: unknown section {question.section!r}."
            )

        if not question.prompt:
            errors.append(f"{question.qid}: missing Prompt.")

        if set(question.choices) != set("ABCD"):
            errors.append(
                f"{question.qid}: expected choices A-D, found "
                f"{sorted(question.choices)}."
            )

        for letter in "ABCD":
            if not question.choices.get(letter, "").strip():
                errors.append(f"{question.qid}: empty choice {letter}.")

        normalized_choices = [
            normalize_option_text(question.choices.get(letter, ""))
            for letter in "ABCD"
        ]

        duplicates = [
            text
            for text, count in Counter(normalized_choices).items()
            if text and count > 1
        ]

        if duplicates:
            warnings.append(f"{question.qid}: duplicate option text detected.")

        if question.correct not in "ABCD":
            errors.append(
                f"{question.qid}: missing or invalid correct-answer letter."
            )
        else:
            actual_correct_text = question.choices.get(question.correct, "")

            if (
                question.correct_answer
                and clean_text(question.correct_answer)
                != clean_text(actual_correct_text)
            ):
                warnings.append(
                    f"{question.qid}: Correct Answer text does not exactly "
                    f"match choice {question.correct}."
                )

            question.correct_answer = actual_correct_text

        if not question.explanation:
            errors.append(f"{question.qid}: missing explanation.")

        stimulus_id = question.stimulus_id.strip()

        if stimulus_id.lower() != "none":
            if stimulus_id not in stimuli:
                errors.append(
                    f"{question.qid}: referenced stimulus {stimulus_id} "
                    f"was not found."
                )

    used_stimuli = {
        question.stimulus_id
        for question in questions
        if question.stimulus_id.lower() != "none"
    }

    unused_stimuli = sorted(set(stimuli) - used_stimuli)

    for stimulus_id in unused_stimuli:
        warnings.append(f"Unused stimulus: {stimulus_id}")

    letters = [
        question.correct
        for question in questions
        if question.correct in "ABCD"
    ]

    max_run = 0
    current_run = 0
    previous = None

    for letter in letters:
        if letter == previous:
            current_run += 1
        else:
            current_run = 1
            previous = letter

        max_run = max(max_run, current_run)

    if max_run > 3:
        warnings.append(
            f"Correct-answer sequence contains a run of {max_run} "
            f"identical letters."
        )

    if len(letters) >= 16:
        for start in range(len(letters) - 15):
            block = "".join(letters[start:start + 16])

            if block in {
                "ABCDABCDABCDABCD",
                "BCDABCDABCDABCDA",
                "CDABCDABCDABCDAB",
                "DABCDABCDABCDABC",
            }:
                warnings.append(
                    f"Obvious ABCD cycle detected near question "
                    f"{start + 1}."
                )
                break

    return errors, warnings


def question_to_json(
    question: Question,
    stimuli: dict[str, Stimulus],
) -> dict[str, Any]:
    result: dict[str, Any] = {
        "id": question.qid,
        "itemType": "mcq_single",
        "section": question.section,
        "category": question.category or question.section,
        "skill": question.skill,
        "prompt": question.prompt,
        "choices": {
            letter: question.choices[letter]
            for letter in "ABCD"
        },
        "correct": question.correct,
        "correctAnswer": question.choices[question.correct],
        "explanation": question.explanation,
    }

    stimulus_id = question.stimulus_id.strip()

    if stimulus_id.lower() != "none":
        stimulus = stimuli[stimulus_id]

        result["stimulusId"] = stimulus.stimulus_id
        result["stimulusType"] = stimulus.stimulus_type
        result["stimulusTitle"] = stimulus.title
        result["stimulusText"] = stimulus.text

    return result


def write_exam_json(
    family: str,
    exam_no: int,
    questions: list[Question],
    stimuli: dict[str, Stimulus],
) -> list[Path]:
    written: list[Path] = []
    config = FAMILY_CONFIG[family]

    for section in config["sections"]:
        section_questions = [
            question_to_json(question, stimuli)
            for question in questions
            if question.section == section["source"]
        ]

        filename = (
            f'{family}_{section["slug"]}_exam_{exam_no:02d}.json'
        )

        destination = DATA_DIR / filename

        payload = {
            "title": (
                f'{config["display"]} {section["source"]} '
                f'Practice Test {exam_no:02d}'
            ),
            "examFamily": family,
            "examNumber": exam_no,
            "section": section["source"],
            "questionCount": len(section_questions),
            "questions": section_questions,
        }

        destination.write_text(
            json.dumps(payload, indent=2, ensure_ascii=False) + "\n",
            encoding="utf-8",
        )

        written.append(destination)

    return written


def process_exam(
    family: str,
    exam_no: int,
    parts: dict[str, Path],
    dry_run: bool,
) -> tuple[bool, list[str]]:
    messages: list[str] = []
    questions: list[Question] = []
    stimuli: dict[str, Stimulus] = {}
    parse_warnings: list[str] = []

    required_parts = set("abcd")
    present_parts = set(parts)

    missing_parts = sorted(required_parts - present_parts)

    if missing_parts:
        messages.append(
            f"ERROR: Missing parts: {', '.join(part.upper() for part in missing_parts)}"
        )
        return False, messages

    for part in "abcd":
        path = parts[part]
        part_questions, part_stimuli, warnings = parse_part(path)

        messages.append(
            f"{path.name}: {len(part_questions)} questions, "
            f"{len(part_stimuli)} stimuli"
        )

        questions.extend(part_questions)
        stimuli.update(part_stimuli)

        for warning in warnings:
            parse_warnings.append(f"{path.name}: {warning}")

    override_notes = apply_overrides(questions)

    errors, warnings = validate_exam(
        family,
        exam_no,
        questions,
        stimuli,
        parts,
    )

    warnings = parse_warnings + warnings

    report_lines = [
        f"Family: {family.upper()}",
        f"Exam: {exam_no:02d}",
        f"Questions parsed: {len(questions)}",
        f"Stimuli parsed: {len(stimuli)}",
        "",
    ]

    if errors:
        report_lines.append("ERRORS")
        report_lines.extend(f"- {error}" for error in errors)
        report_lines.append("")

    if warnings:
        report_lines.append("WARNINGS")
        report_lines.extend(f"- {warning}" for warning in warnings)
        report_lines.append("")

    if not errors:
        if dry_run:
            report_lines.append("RESULT")
            report_lines.append("- Validation passed; no JSON written.")
        else:
            written = write_exam_json(
                family,
                exam_no,
                questions,
                stimuli,
            )

            report_lines.append("RESULT")
            report_lines.append("- Validation passed.")
            report_lines.append(f"- JSON files written: {len(written)}")

            for path in written:
                report_lines.append(f"  {path.relative_to(ROOT)}")

    report_name = f"{family}_exam_{exam_no:02d}_report.txt"
    report_path = REPORTS_DIR / report_name

    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    report_path.write_text(
        "\n".join(report_lines).rstrip() + "\n",
        encoding="utf-8",
    )

    messages.extend(report_lines)
    messages.append(f"Report: {report_path.relative_to(ROOT)}")

    return not errors, messages


def validate_generated_json() -> list[str]:
    errors: list[str] = []

    for family, config in FAMILY_CONFIG.items():
        for section in config["sections"]:
            for exam_no in range(1, 11):
                filename = (
                    f'{family}_{section["slug"]}_exam_{exam_no:02d}.json'
                )

                path = DATA_DIR / filename

                if not path.exists():
                    continue

                try:
                    payload = json.loads(path.read_text(encoding="utf-8"))
                except Exception as exc:
                    errors.append(f"{filename}: invalid JSON: {exc}")
                    continue

                questions = payload.get("questions", [])

                if len(questions) != section["count"]:
                    errors.append(
                        f"{filename}: expected {section['count']} questions, "
                        f"found {len(questions)}."
                    )

                ids = [question.get("id") for question in questions]

                if len(ids) != len(set(ids)):
                    errors.append(f"{filename}: duplicate question IDs.")

    return errors


def main() -> int:
    parser = argparse.ArgumentParser(
        description=(
            "Import four-part HSPT, TACHS, and COOP practice exams "
            "from DOCX or TXT into platform JSON."
        )
    )

    parser.add_argument(
        "--family",
        choices=["hspt", "tachs", "coop", "all"],
        default="all",
    )

    parser.add_argument(
        "--exam",
        type=int,
        help="Import only a specific exam number.",
    )

    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="Validate without writing JSON files.",
    )

    parser.add_argument(
        "--validate-json",
        action="store_true",
        help="Validate existing generated JSON files.",
    )

    args = parser.parse_args()

    DATA_DIR.mkdir(parents=True, exist_ok=True)
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)

    if args.validate_json:
        errors = validate_generated_json()

        if errors:
            print("JSON VALIDATION FAILED")

            for error in errors:
                print("ERROR:", error)

            return 1

        print("JSON VALIDATION PASSED")
        return 0

    families = (
        list(FAMILY_CONFIG)
        if args.family == "all"
        else [args.family]
    )

    overall_success = True
    processed = 0

    for family in families:
        discovered = discover_sources(family)

        if args.exam is not None:
            discovered = {
                exam_no: parts
                for exam_no, parts in discovered.items()
                if exam_no == args.exam
            }

        if not discovered:
            print(f"{family.upper()}: no matching exam files found.")
            continue

        for exam_no, parts in sorted(discovered.items()):
            processed += 1

            print()
            print("=" * 64)
            print(f"{family.upper()} EXAM {exam_no:02d}")
            print("=" * 64)

            success, messages = process_exam(
                family,
                exam_no,
                parts,
                args.dry_run,
            )

            for message in messages:
                print(message)

            overall_success = overall_success and success

    if processed == 0:
        print()
        print("No complete exam sources were found.")
        print()
        print("Expected examples:")
        print("imports/hspt/hspt_full_exam_01_part_a.docx")
        print("imports/hspt/hspt_full_exam_01_part_b.docx")
        print("imports/hspt/hspt_full_exam_01_part_c.docx")
        print("imports/hspt/hspt_full_exam_01_part_d.docx")
        return 0

    print()
    print("=" * 64)

    if overall_success:
        print("IMPORT COMPLETED SUCCESSFULLY")
    else:
        print("IMPORT COMPLETED WITH BLOCKING ERRORS")

    print("=" * 64)

    return 0 if overall_success else 1


if __name__ == "__main__":
    sys.exit(main())
